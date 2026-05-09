from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ACCENT = RGBColor(22, 89, 47)
ACCENT_DARK = RGBColor(15, 58, 27)
ACCENT_LIGHT = RGBColor(231, 242, 235)
TEXT = RGBColor(33, 37, 41)
MUTED = RGBColor(102, 112, 122)
BORDER = "D9E2D9"

ROOT = Path(__file__).resolve().parents[2]
OUTPUT_DIR = Path(r"D:\Freelance\DestinyDevelopment\Proposals")
OUTPUT_PATH = OUTPUT_DIR / "Hectors_Tree_Service_Website_Marketing_Proposal_Discounted.docx"
LOGO_PATH = Path(r"C:\Users\Dady1\Downloads\HLogo_New.png")


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, color=BORDER, size="8") -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "start", "bottom", "end"):
        el = tc_borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            tc_borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), size)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)


def set_table_layout(table) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_end])


def apply_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    normal.font.color.rgb = TEXT
    pf = normal.paragraph_format
    pf.line_spacing = 1.08
    pf.space_after = Pt(6)

    title = doc.styles["Title"]
    title.font.name = "Arial"
    title.font.size = Pt(24)
    title.font.bold = True
    title.font.color.rgb = ACCENT_DARK
    title.paragraph_format.space_after = Pt(4)

    subtitle = doc.styles["Subtitle"]
    subtitle.font.name = "Arial"
    subtitle.font.size = Pt(12)
    subtitle.font.color.rgb = MUTED
    subtitle.paragraph_format.space_after = Pt(12)

    for style_name, size in (("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11)):
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = ACCENT_DARK
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(4)


def setup_page(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("DestinyDevelopment | Website & Growth Proposal")
    run.font.name = "Arial"
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED
    p.paragraph_format.space_after = Pt(4)

    border_p = header.add_paragraph()
    border_p.paragraph_format.space_after = Pt(0)
    p_pr = border_p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "D9E2D9")
    pbdr.append(bottom)
    p_pr.append(pbdr)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fr = fp.add_run("Page ")
    fr.font.name = "Arial"
    fr.font.size = Pt(9)
    fr.font.color.rgb = MUTED
    add_page_number(fp)


def add_cover(doc: Document) -> None:
    if LOGO_PATH.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(LOGO_PATH), width=Inches(2.3))

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Website Creation, Marketing & Maintenance Proposal")

    sub = doc.add_paragraph(style="Subtitle")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run("Prepared for Hector's Tree Service")

    meta = doc.add_table(rows=4, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta.autofit = False
    set_table_layout(meta)
    meta.columns[0].width = Inches(2.2)
    meta.columns[1].width = Inches(3.8)
    items = [
        ("Client", "Hector Aguilar"),
        ("Business", "Hector's Tree Service"),
        ("Prepared by", "Israel Aguilar, CEO | DestinyDevelopment"),
        ("Date", date(2026, 5, 9).strftime("%B %d, %Y")),
    ]
    for row, (label, value) in zip(meta.rows, items):
        for cell in row.cells:
            set_cell_border(cell)
            set_cell_margins(cell, top=80, bottom=80)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        left, right = row.cells
        set_cell_shading(left, "E7F2EB")
        lp = left.paragraphs[0]
        lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        lrun = lp.add_run(label)
        lrun.font.bold = True
        lrun.font.color.rgb = ACCENT_DARK
        rp = right.paragraphs[0]
        rr = rp.add_run(value)
        rr.font.color.rgb = TEXT

    doc.add_paragraph()
    callout = doc.add_table(rows=1, cols=1)
    callout.alignment = WD_TABLE_ALIGNMENT.CENTER
    callout.autofit = False
    set_table_layout(callout)
    callout.columns[0].width = Inches(6.2)
    cell = callout.cell(0, 0)
    set_cell_shading(cell, "F5F9F6")
    set_cell_border(cell, color="B7CFBC", size="10")
    set_cell_margins(cell, top=150, bottom=150, start=180, end=180)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(
        "Objective: deliver a professional lead-generating website, dependable monthly maintenance, "
        "and consistent content marketing that helps Hector's Tree Service win higher-value jobs "
        "without relying heavily on pay-per-lead platforms."
    )
    r.font.size = Pt(11)
    r.font.color.rgb = TEXT
    doc.add_paragraph()
    promo = doc.add_table(rows=1, cols=1)
    promo.alignment = WD_TABLE_ALIGNMENT.CENTER
    promo.autofit = False
    set_table_layout(promo)
    promo.columns[0].width = Inches(6.2)
    promo_cell = promo.cell(0, 0)
    set_cell_shading(promo_cell, "16592F")
    set_cell_border(promo_cell, color="16592F", size="10")
    set_cell_margins(promo_cell, top=160, bottom=160, start=180, end=180)
    promo_p = promo_cell.paragraphs[0]
    promo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    promo_r1 = promo_p.add_run("Limited-Time Website Launch Offer: 25% Off the Website Build")
    promo_r1.font.bold = True
    promo_r1.font.size = Pt(13)
    promo_r1.font.color.rgb = RGBColor(255, 255, 255)
    promo_p2 = promo_cell.add_paragraph()
    promo_p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    promo_r2 = promo_p2.add_run("Standard build value: $3,300 | Discounted launch investment: $2,475")
    promo_r2.font.size = Pt(11)
    promo_r2.font.color.rgb = RGBColor(255, 255, 255)
    doc.add_page_break()


def add_intro(doc: Document) -> None:
    doc.add_paragraph("Executive Summary", style="Heading 1")
    for text in [
        "This proposal combines three things local service companies usually have to buy separately: a custom website, ongoing site care, and monthly content marketing. By keeping those pieces together, the business gets a cleaner brand experience, faster execution, and one accountable partner.",
        "For a tree service company in the Nashville market, the website must do more than look good. It needs to earn trust quickly, convert mobile visitors into estimate requests, and support repeat visibility through Google, social media, and review-driven content.",
        "The pricing below is structured to be fair for the Nashville, TN market while still reflecting the real production work involved in design, development, maintenance, on-site content capture, editing, and business workflow support.",
    ]:
        doc.add_paragraph(text)


def add_website_investment(doc: Document) -> None:
    doc.add_paragraph("One-Time Website Creation Investment", style="Heading 1")
    doc.add_paragraph(
        "The website build fee covers planning, design, development, mobile optimization, launch preparation, and revision support for a polished lead-generation site."
    )
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_layout(table)
    widths = [Inches(3.8), Inches(1.15), Inches(1.55)]
    for idx, width in enumerate(widths):
        table.columns[idx].width = width

    headers = ["Scope Item", "Hours", "Value"]
    for cell, text in zip(table.rows[0].cells, headers):
        set_cell_shading(cell, "16592F")
        set_cell_border(cell, color="16592F", size="10")
        set_cell_margins(cell, top=90, bottom=90)
        p = cell.paragraphs[0]
        r = p.add_run(text)
        r.font.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)

    rows = [
        ("Discovery, planning, sitemap, messaging structure", "4", "$320"),
        ("Custom UI/UX design and mobile-first layout refinement", "8", "$760"),
        ("Frontend build, estimate funnel, CTA integration, routing", "12", "$1,260"),
        ("Performance tuning, responsive QA, SEO foundations", "5", "$475"),
        ("Launch support, review rounds, content population, polish", "5", "$485"),
    ]
    for scope, hours, value in rows:
        row = table.add_row()
        for cell in row.cells:
            set_cell_border(cell)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        row.cells[0].paragraphs[0].add_run(scope)
        row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row.cells[1].paragraphs[0].add_run(hours)
        row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        row.cells[2].paragraphs[0].add_run(value)

    total = table.add_row()
    for cell in total.cells:
        set_cell_border(cell, color="8FAF95", size="10")
        set_cell_margins(cell, top=100, bottom=100)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_shading(cell, "F5F9F6")
    total.cells[0].paragraphs[0].add_run("Total Website Creation Investment").font.bold = True
    total.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    total.cells[1].paragraphs[0].add_run("34").font.bold = True
    total.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    tr = total.cells[2].paragraphs[0].add_run("$3,300")
    tr.font.bold = True
    tr.font.color.rgb = ACCENT_DARK

    discount = table.add_row()
    for cell in discount.cells:
        set_cell_border(cell, color="16592F", size="10")
        set_cell_margins(cell, top=110, bottom=110)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_shading(cell, "E7F2EB")
    discount.cells[0].paragraphs[0].add_run("25% Launch Discount Applied").font.bold = True
    discount.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    discount.cells[1].paragraphs[0].add_run("-").font.bold = True
    discount.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    dr = discount.cells[2].paragraphs[0].add_run("-$825")
    dr.font.bold = True
    dr.font.color.rgb = ACCENT_DARK

    final_total = table.add_row()
    for cell in final_total.cells:
        set_cell_border(cell, color="16592F", size="12")
        set_cell_margins(cell, top=120, bottom=120)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_shading(cell, "DDEEE1")
    final_total.cells[0].paragraphs[0].add_run("Discounted Website Investment").font.bold = True
    final_total.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    final_total.cells[1].paragraphs[0].add_run("34").font.bold = True
    final_total.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fr = final_total.cells[2].paragraphs[0].add_run("$2,475")
    fr.font.bold = True
    fr.font.color.rgb = ACCENT_DARK

    note = doc.add_paragraph()
    note.add_run("Recommended billing structure with discount: ").font.bold = True
    note.add_run("$1,250 to initiate the project and $1,225 upon final launch approval.")


def add_package_summary(doc: Document) -> None:
    doc.add_paragraph("Monthly Growth Packages", style="Heading 1")
    doc.add_paragraph(
        "All packages include website hosting coordination, routine maintenance, security checks, and early access to the DocSetter mobile app. The difference is the monthly content volume, marketing support level, and depth of operational help."
    )

    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_layout(table)
    for idx, width in enumerate([Inches(2.2), Inches(1.43), Inches(1.43), Inches(1.44)]):
        table.columns[idx].width = width

    head = ["Feature", "Starter", "Growth", "Market Leader"]
    for cell, text in zip(table.rows[0].cells, head):
        set_cell_shading(cell, "16592F")
        set_cell_border(cell, color="16592F", size="10")
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        r = p.add_run(text)
        r.font.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)

    summary_rows = [
        ("Monthly investment", "$595", "$995", "$1,595"),
        ("Branded posts per month", "8", "12", "16"),
        ("Edited short-form videos", "2", "4", "8"),
        ("Drone video sessions", "1 mini shoot", "1 full shoot", "2 shoots"),
        ("Website update requests", "1", "2", "Priority ongoing"),
        ("Reporting cadence", "Monthly snapshot", "Monthly review", "Monthly strategy + review"),
        ("DocSetter early access", "1 seat", "2 seats", "Up to 5 seats"),
    ]

    for feature, starter, growth, leader in summary_rows:
        row = table.add_row()
        for i, cell in enumerate(row.cells):
            set_cell_border(cell)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if i == 0:
                cell.paragraphs[0].add_run(feature).font.bold = True
            else:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row.cells[1].paragraphs[0].add_run(starter)
        row.cells[2].paragraphs[0].add_run(growth)
        row.cells[3].paragraphs[0].add_run(leader)


def add_package_cards(doc: Document) -> None:
    packages = [
        (
            "Starter Package | $595 / month",
            "Best for staying active online without taking on a heavy monthly retainer.",
            [
                "8 branded social posts each month",
                "2 edited short-form videos each month",
                "1 monthly drone mini-session for fresh aerial content",
                "1 routine website content/update request each month",
                "Hosting coordination, backups, uptime check, plugin/theme updates, and form monitoring",
                "Monthly analytics snapshot and basic recommendations",
                "Early access to DocSetter with 1 user seat",
            ],
        ),
        (
            "Growth Package | $995 / month",
            "Recommended for local lead generation, stronger social consistency, and steady website improvement.",
            [
                "12 branded social posts each month",
                "4 edited short-form videos each month",
                "1 full drone content session each month",
                "2 website update requests each month plus seasonal content refreshes",
                "Google Business Profile content support and review-request asset support",
                "Monthly reporting review and one strategy check-in",
                "Early access to DocSetter with 2 user seats and onboarding support",
            ],
        ),
        (
            "Market Leader Package | $1,595 / month",
            "Best for scaling visibility faster and turning content into a stronger sales asset.",
            [
                "16 branded social posts each month",
                "8 edited short-form videos/reels each month",
                "2 drone production sessions each month",
                "Priority website edits, landing-page support, and monthly offer/campaign adjustments",
                "Monthly reporting review, strategy call, and content calendar planning",
                "Review-growth support and deeper brand consistency across channels",
                "Early access to DocSetter with up to 5 seats and workflow setup guidance",
            ],
        ),
    ]

    for title, subtitle, bullets in packages:
        doc.add_paragraph(title, style="Heading 2")
        blurb = doc.add_paragraph()
        br = blurb.add_run(subtitle)
        br.italic = True
        br.font.color.rgb = MUTED
        for bullet in bullets:
            p = doc.add_paragraph(style="Normal")
            p.paragraph_format.left_indent = Inches(0.35)
            p.paragraph_format.first_line_indent = Inches(-0.2)
            run = p.add_run("• ")
            run.font.bold = True
            p.add_run(bullet)


def add_roi(doc: Document) -> None:
    doc.add_paragraph("Why This Investment Makes Business Sense", style="Heading 1")
    doc.add_paragraph(
        "Local service companies usually compare monthly retainers to what they would spend on leads, admin software, and time lost chasing follow-ups. That is the right lens."
    )

    roi = doc.add_table(rows=1, cols=2)
    roi.alignment = WD_TABLE_ALIGNMENT.LEFT
    roi.autofit = False
    set_table_layout(roi)
    roi.columns[0].width = Inches(3.15)
    roi.columns[1].width = Inches(3.35)

    cards = [
        (
            "Lead Cost Comparison",
            [
                "Recent contractor lead guides report shared-lead costs often landing around $15 to $80 per lead, with customer acquisition climbing much higher once close rates are considered.",
                "At only 15 paid leads per month, that can mean roughly $225 to $1,200 in lead fees before the job is even sold.",
                "A branded website and content system builds an owned channel instead of renting attention one lead at a time.",
            ],
        ),
        (
            "Time Savings With DocSetter",
            [
                "DocSetter is positioned to reduce administrative drag around project organization, routing, and field coordination.",
                "If the owner or office team saves even 8 to 15 hours per month, and that time is worth $60 to $90 per hour, the recovered value is roughly $480 to $1,350 per month.",
                "That time can be redirected into estimates, scheduling, upsells, and faster response time to inbound jobs.",
            ],
        ),
    ]

    for cell, (heading, bullets) in zip(roi.rows[0].cells, cards):
        set_cell_shading(cell, "F5F9F6")
        set_cell_border(cell, color="B7CFBC", size="10")
        set_cell_margins(cell, top=140, bottom=140, start=160, end=160)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        hp = cell.paragraphs[0]
        hr = hp.add_run(heading)
        hr.font.bold = True
        hr.font.color.rgb = ACCENT_DARK
        for bullet in bullets:
            p = cell.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.22)
            p.paragraph_format.first_line_indent = Inches(-0.16)
            p.add_run("• ").bold = True
            p.add_run(bullet)

    doc.add_paragraph()
    note = doc.add_paragraph()
    note.add_run("Reference point: ").font.bold = True
    note.add_run(
        "even lightweight contractor admin apps can carry their own monthly subscription cost, "
        "while fuller routing and field-service platforms add more spend on top. "
        "Including early DocSetter access inside these packages increases the practical value without adding another separate invoice right away."
    )


def add_recommendation(doc: Document) -> None:
    doc.add_paragraph("Recommended Path", style="Heading 1")
    doc.add_paragraph(
        "For Hector's Tree Service, the recommended option is the Growth Package at $995 per month. "
        "It is the strongest balance of visible marketing output, website care, and operational value without overcommitting too early."
    )
    doc.add_paragraph(
        "That package is large enough to keep the brand active with consistent posts and videos, while still staying within a reasonable budget for a local service company. "
        "It also gives the business enough content to strengthen Google presence, social proof, and repeat visibility over time."
    )

    doc.add_paragraph("Approval & Next Steps", style="Heading 1")
    doc.add_paragraph(
        "If approved, the project can move forward with website launch support, monthly content scheduling, and DocSetter onboarding planning. "
        "This proposal is valid for 30 days from the date above."
    )

    doc.add_paragraph("Helpful Add-Ons", style="Heading 2")
    add_on_table = doc.add_table(rows=1, cols=2)
    add_on_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    add_on_table.autofit = False
    set_table_layout(add_on_table)
    add_on_table.columns[0].width = Inches(4.4)
    add_on_table.columns[1].width = Inches(1.8)
    for cell, text in zip(add_on_table.rows[0].cells, ["Optional Add-On", "Rate"]):
        set_cell_shading(cell, "E7F2EB")
        set_cell_border(cell, color="B7CFBC", size="10")
        set_cell_margins(cell)
        run = cell.paragraphs[0].add_run(text)
        run.font.bold = True
        run.font.color.rgb = ACCENT_DARK
    for name, rate in [
        ("Extra edited short-form video", "$125 each"),
        ("Extra drone content session", "$250 each"),
        ("Additional landing page or service page", "$225 each"),
    ]:
        row = add_on_table.add_row()
        for cell in row.cells:
            set_cell_border(cell)
            set_cell_margins(cell)
        row.cells[0].paragraphs[0].add_run(name)
        row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        row.cells[1].paragraphs[0].add_run(rate)

    terms = doc.add_paragraph()
    terms.add_run("Terms note: ").font.bold = True
    terms.add_run(
        "monthly packages are intended to keep output consistent. On-site filming and drone scheduling remain subject to weather, property access, and flight safety conditions."
    )

    prepared = doc.add_paragraph()
    prepared.add_run("Prepared by: ").font.bold = True
    prepared.add_run("Israel Aguilar | CEO, DestinyDevelopment")

    sign = doc.add_table(rows=2, cols=2)
    sign.alignment = WD_TABLE_ALIGNMENT.LEFT
    sign.autofit = False
    set_table_layout(sign)
    sign.columns[0].width = Inches(3.1)
    sign.columns[1].width = Inches(3.1)
    labels = [("Client Approval", "Date"), ("Hector Aguilar", "__________________________")]
    for row_idx, (left_text, right_text) in enumerate(labels):
        row = sign.rows[row_idx]
        for cell in row.cells:
            set_cell_border(cell, color="FFFFFF", size="0")
            set_cell_margins(cell, top=110, bottom=110)
        row.cells[0].paragraphs[0].add_run(left_text).font.bold = row_idx == 0
        row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        row.cells[1].paragraphs[0].add_run(right_text).font.bold = row_idx == 0


def build_document() -> Path:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    apply_styles(doc)
    setup_page(doc)
    add_cover(doc)
    add_intro(doc)
    add_website_investment(doc)
    add_package_summary(doc)
    add_package_cards(doc)
    add_roi(doc)
    add_recommendation(doc)
    doc.save(OUTPUT_PATH)
    return OUTPUT_PATH


if __name__ == "__main__":
    path = build_document()
    print(path)
