from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ACCENT = RGBColor(22, 89, 47)
ACCENT_DARK = RGBColor(15, 58, 27)
TEXT = RGBColor(33, 37, 41)
MUTED = RGBColor(102, 112, 122)
BORDER = "D9E2D9"

OUTPUT_DIR = Path(r"D:\Freelance\DestinyDevelopment\Proposals")
OUTPUT_PATH = OUTPUT_DIR / "Hectors_Tree_Service_Website_Marketing_Proposal_ES.docx"
LOGO_PATH = Path(r"C:\Users\Dady1\Downloads\HLogo_New.png")


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, color=BORDER, size="8") -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "start", "bottom", "end"):
        el = tc_borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            tc_borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), size)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)


def set_table_layout(table) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_end])


def apply_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    normal.font.color.rgb = TEXT
    normal.paragraph_format.line_spacing = 1.08
    normal.paragraph_format.space_after = Pt(6)

    title = doc.styles["Title"]
    title.font.name = "Arial"
    title.font.size = Pt(24)
    title.font.bold = True
    title.font.color.rgb = ACCENT_DARK
    title.paragraph_format.space_after = Pt(4)

    subtitle = doc.styles["Subtitle"]
    subtitle.font.name = "Arial"
    subtitle.font.size = Pt(12)
    subtitle.font.color.rgb = MUTED
    subtitle.paragraph_format.space_after = Pt(12)

    for style_name, size in (("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11)):
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = ACCENT_DARK
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(4)


def setup_page(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("DestinyDevelopment | Propuesta Web y Crecimiento")
    run.font.name = "Arial"
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED

    border_p = header.add_paragraph()
    p_pr = border_p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "D9E2D9")
    pbdr.append(bottom)
    p_pr.append(pbdr)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fr = fp.add_run("Página ")
    fr.font.name = "Arial"
    fr.font.size = Pt(9)
    fr.font.color.rgb = MUTED
    add_page_number(fp)


def add_cover(doc: Document) -> None:
    if LOGO_PATH.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(LOGO_PATH), width=Inches(2.3))

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Propuesta de Creación de Sitio Web, Marketing y Mantenimiento")

    sub = doc.add_paragraph(style="Subtitle")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run("Preparada para Hector's Tree Service")

    meta = doc.add_table(rows=4, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta.autofit = False
    set_table_layout(meta)
    meta.columns[0].width = Inches(2.2)
    meta.columns[1].width = Inches(3.8)
    items = [
        ("Cliente", "Hector Aguilar"),
        ("Negocio", "Hector's Tree Service"),
        ("Preparado por", "Israel Aguilar, CEO | DestinyDevelopment"),
        ("Fecha", date(2026, 5, 9).strftime("%d de %B de %Y")),
    ]
    for row, (label, value) in zip(meta.rows, items):
        for cell in row.cells:
            set_cell_border(cell)
            set_cell_margins(cell, top=80, bottom=80)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        left, right = row.cells
        set_cell_shading(left, "E7F2EB")
        lr = left.paragraphs[0].add_run(label)
        lr.font.bold = True
        lr.font.color.rgb = ACCENT_DARK
        right.paragraphs[0].add_run(value)

    doc.add_paragraph()
    callout = doc.add_table(rows=1, cols=1)
    callout.alignment = WD_TABLE_ALIGNMENT.CENTER
    callout.autofit = False
    set_table_layout(callout)
    callout.columns[0].width = Inches(6.2)
    cell = callout.cell(0, 0)
    set_cell_shading(cell, "F5F9F6")
    set_cell_border(cell, color="B7CFBC", size="10")
    set_cell_margins(cell, top=150, bottom=150, start=180, end=180)
    cell.paragraphs[0].add_run(
        "Objetivo: entregar un sitio web profesional que genere prospectos, mantenimiento mensual confiable "
        "y marketing de contenido constante que ayude a Hector's Tree Service a conseguir trabajos de mayor valor "
        "sin depender demasiado de plataformas de leads pagados."
    )

    doc.add_paragraph()
    promo = doc.add_table(rows=1, cols=1)
    promo.alignment = WD_TABLE_ALIGNMENT.CENTER
    promo.autofit = False
    set_table_layout(promo)
    promo.columns[0].width = Inches(6.2)
    promo_cell = promo.cell(0, 0)
    set_cell_shading(promo_cell, "16592F")
    set_cell_border(promo_cell, color="16592F", size="10")
    set_cell_margins(promo_cell, top=160, bottom=160, start=180, end=180)
    p1 = promo_cell.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p1.add_run("Oferta de Lanzamiento por Tiempo Limitado: 25% de Descuento en la Creación del Sitio Web")
    r1.font.bold = True
    r1.font.size = Pt(13)
    r1.font.color.rgb = RGBColor(255, 255, 255)
    p2 = promo_cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Valor estándar del proyecto: $3,300 | Inversión con descuento: $2,475")
    r2.font.size = Pt(11)
    r2.font.color.rgb = RGBColor(255, 255, 255)


def add_body(doc: Document) -> None:
    doc.add_paragraph("Resumen Ejecutivo", style="Heading 1")
    for text in [
        "Esta propuesta reúne tres servicios que normalmente una empresa local tendría que contratar por separado: un sitio web personalizado, mantenimiento continuo y marketing de contenido mensual. Al mantener todo en un solo servicio, el negocio obtiene una experiencia de marca más sólida, ejecución más rápida y un solo responsable del resultado.",
        "Para una empresa de árboles en el mercado de Nashville, el sitio web no solo debe verse bien. Debe generar confianza rápidamente, convertir visitantes móviles en solicitudes de estimate y respaldar la visibilidad constante a través de Google, redes sociales y contenido basado en reputación.",
        "La inversión presentada a continuación está estructurada para ser justa en el mercado de Nashville, TN, reflejando el trabajo real que implica el diseño, desarrollo, mantenimiento, producción de contenido, edición y apoyo operativo.",
    ]:
        doc.add_paragraph(text)

    doc.add_paragraph("Inversión Única para la Creación del Sitio Web", style="Heading 1")
    doc.add_paragraph(
        "La inversión del sitio web cubre planeación, diseño, desarrollo, optimización móvil, preparación para lanzamiento y soporte de revisiones para entregar una presencia digital profesional y orientada a generar clientes."
    )

    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_layout(table)
    widths = [Inches(3.8), Inches(1.15), Inches(1.55)]
    for idx, width in enumerate(widths):
        table.columns[idx].width = width

    for cell, text in zip(table.rows[0].cells, ["Concepto", "Horas", "Valor"]):
        set_cell_shading(cell, "16592F")
        set_cell_border(cell, color="16592F", size="10")
        set_cell_margins(cell)
        run = cell.paragraphs[0].add_run(text)
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)

    rows = [
        ("Descubrimiento, planeación, mapa del sitio y estructura de mensajes", "4", "$320"),
        ("Diseño UI/UX personalizado y refinamiento mobile-first", "8", "$760"),
        ("Desarrollo frontend, embudo de estimate, CTAs e integración de rutas", "12", "$1,260"),
        ("Optimización de rendimiento, QA responsivo y bases SEO", "5", "$475"),
        ("Soporte de lanzamiento, revisiones, carga de contenido y pulido final", "5", "$485"),
    ]
    for scope, hours, value in rows:
        row = table.add_row()
        for cell in row.cells:
            set_cell_border(cell)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        row.cells[0].paragraphs[0].add_run(scope)
        row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row.cells[1].paragraphs[0].add_run(hours)
        row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        row.cells[2].paragraphs[0].add_run(value)

    totals = [
        ("Inversión Total del Sitio Web", "34", "$3,300", "F5F9F6"),
        ("Descuento de Lanzamiento 25% Aplicado", "-", "-$825", "E7F2EB"),
        ("Inversión Final con Descuento", "34", "$2,475", "DDEEE1"),
    ]
    for label, hours, value, fill in totals:
        row = table.add_row()
        for cell in row.cells:
            set_cell_border(cell, color="16592F" if fill != "F5F9F6" else "8FAF95", size="10")
            set_cell_margins(cell, top=110, bottom=110)
            set_cell_shading(cell, fill)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        row.cells[0].paragraphs[0].add_run(label).font.bold = True
        row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row.cells[1].paragraphs[0].add_run(hours).font.bold = True
        row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        rr = row.cells[2].paragraphs[0].add_run(value)
        rr.font.bold = True
        rr.font.color.rgb = ACCENT_DARK

    pay = doc.add_paragraph()
    pay.add_run("Forma de pago sugerida con descuento: ").font.bold = True
    pay.add_run("$1,250 para iniciar el proyecto y $1,225 al aprobar el lanzamiento final.")

    doc.add_paragraph("Paquetes Mensuales de Crecimiento", style="Heading 1")
    doc.add_paragraph(
        "Todos los paquetes incluyen coordinación de hosting, mantenimiento del sitio, revisiones básicas de seguridad y acceso anticipado a la aplicación móvil DocSetter. La diferencia principal está en el volumen mensual de contenido, el nivel de marketing y la profundidad del apoyo operativo."
    )

    summary = doc.add_table(rows=1, cols=4)
    summary.alignment = WD_TABLE_ALIGNMENT.LEFT
    summary.autofit = False
    set_table_layout(summary)
    for idx, width in enumerate([Inches(2.2), Inches(1.43), Inches(1.43), Inches(1.44)]):
        summary.columns[idx].width = width
    for cell, text in zip(summary.rows[0].cells, ["Característica", "Starter", "Growth", "Market Leader"]):
        set_cell_shading(cell, "16592F")
        set_cell_border(cell, color="16592F", size="10")
        set_cell_margins(cell)
        run = cell.paragraphs[0].add_run(text)
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)

    summary_rows = [
        ("Inversión mensual", "$595", "$995", "$1,595"),
        ("Posts de marca por mes", "8", "12", "16"),
        ("Videos cortos editados", "2", "4", "8"),
        ("Sesiones con dron", "1 sesión mini", "1 sesión completa", "2 sesiones"),
        ("Solicitudes de actualización web", "1", "2", "Prioridad continua"),
        ("Reportes", "Resumen mensual", "Revisión mensual", "Estrategia + revisión mensual"),
        ("Acceso temprano a DocSetter", "1 usuario", "2 usuarios", "Hasta 5 usuarios"),
    ]
    for feature, starter, growth, leader in summary_rows:
        row = summary.add_row()
        for i, cell in enumerate(row.cells):
            set_cell_border(cell)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if i == 0:
                cell.paragraphs[0].add_run(feature).font.bold = True
            else:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row.cells[1].paragraphs[0].add_run(starter)
        row.cells[2].paragraphs[0].add_run(growth)
        row.cells[3].paragraphs[0].add_run(leader)

    packages = [
        (
            "Starter | $595 / mes",
            "Ideal para mantener presencia constante sin asumir un retainer mensual demasiado alto.",
            [
                "8 posts de marca al mes",
                "2 videos cortos editados al mes",
                "1 sesión mini mensual con dron para contenido aéreo fresco",
                "1 solicitud mensual de actualización web",
                "Coordinación de hosting, respaldos, revisión de uptime, actualizaciones y monitoreo del formulario",
                "Resumen mensual de analítica con recomendaciones básicas",
                "Acceso temprano a DocSetter con 1 usuario",
            ],
        ),
        (
            "Growth | $995 / mes",
            "Recomendado para generación local de clientes, mejor consistencia en redes y evolución constante del sitio web.",
            [
                "12 posts de marca al mes",
                "4 videos cortos editados al mes",
                "1 sesión completa de contenido con dron al mes",
                "2 solicitudes mensuales de actualización web y refresco de contenido estacional",
                "Apoyo de contenido para Google Business Profile y activos para pedir reviews",
                "Revisión mensual de reportes y una sesión estratégica",
                "Acceso temprano a DocSetter con 2 usuarios y apoyo inicial de uso",
            ],
        ),
        (
            "Market Leader | $1,595 / mes",
            "Ideal para escalar visibilidad más rápido y convertir el contenido en un activo comercial más fuerte.",
            [
                "16 posts de marca al mes",
                "8 videos cortos o reels editados al mes",
                "2 sesiones mensuales de producción con dron",
                "Ediciones web prioritarias, apoyo para landing pages y ajustes mensuales de campañas",
                "Revisión mensual de reportes, llamada estratégica y planeación de calendario de contenido",
                "Apoyo para crecimiento de reviews y mayor consistencia de marca",
                "Acceso temprano a DocSetter con hasta 5 usuarios y guía de configuración de flujo",
            ],
        ),
    ]
    for title, subtitle, bullets in packages:
        doc.add_paragraph(title, style="Heading 2")
        sp = doc.add_paragraph()
        sr = sp.add_run(subtitle)
        sr.italic = True
        sr.font.color.rgb = MUTED
        for bullet in bullets:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.35)
            p.paragraph_format.first_line_indent = Inches(-0.2)
            p.add_run("• ").bold = True
            p.add_run(bullet)

    doc.add_paragraph("Por Qué Esta Inversión Tiene Sentido", style="Heading 1")
    doc.add_paragraph(
        "A los negocios de servicios les interesa el resultado económico. Por eso esta propuesta está diseñada para compararse contra lo que normalmente se pierde en leads pagados, tiempo administrativo y procesos manuales."
    )

    roi = doc.add_table(rows=1, cols=2)
    roi.alignment = WD_TABLE_ALIGNMENT.LEFT
    roi.autofit = False
    set_table_layout(roi)
    roi.columns[0].width = Inches(3.15)
    roi.columns[1].width = Inches(3.35)
    cards = [
        (
            "Comparación Contra Leads Pagados",
            [
                "Guías recientes para contratistas muestran costos compartidos por lead que muchas veces caen entre $15 y $80 por lead, y el costo real sube todavía más cuando se analiza la tasa de cierre.",
                "Con solo 15 leads pagados al mes, un negocio puede terminar gastando aproximadamente entre $225 y $1,200 antes siquiera de vender el trabajo.",
                "Un sitio web propio y una estrategia de contenido ayudan a construir un canal propio, en lugar de seguir rentando atención lead por lead.",
            ],
        ),
        (
            "Ahorro de Tiempo con DocSetter",
            [
                "DocSetter está pensado para reducir carga administrativa en organización de proyectos, rutas y coordinación del trabajo.",
                "Si el dueño o el equipo ahorran entre 8 y 15 horas por mes, y ese tiempo vale entre $60 y $90 por hora, el valor recuperado puede rondar entre $480 y $1,350 mensuales.",
                "Ese tiempo recuperado se puede redirigir a estimates, programación, upsells y respuesta más rápida a nuevos clientes.",
            ],
        ),
    ]
    for cell, (heading, bullets) in zip(roi.rows[0].cells, cards):
        set_cell_shading(cell, "F5F9F6")
        set_cell_border(cell, color="B7CFBC", size="10")
        set_cell_margins(cell, top=140, bottom=140, start=160, end=160)
        head = cell.paragraphs[0].add_run(heading)
        head.font.bold = True
        head.font.color.rgb = ACCENT_DARK
        for bullet in bullets:
            p = cell.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.22)
            p.paragraph_format.first_line_indent = Inches(-0.16)
            p.add_run("• ").bold = True
            p.add_run(bullet)

    ref = doc.add_paragraph()
    ref.add_run("Punto de referencia: ").font.bold = True
    ref.add_run(
        "incluso aplicaciones administrativas básicas para contratistas suelen tener su propio costo mensual, mientras plataformas más completas de rutas y field service agregan todavía más gasto. Incluir acceso temprano a DocSetter dentro de estos paquetes aumenta el valor práctico sin añadir otra factura separada por ahora."
    )

    doc.add_paragraph("Ruta Recomendada", style="Heading 1")
    doc.add_paragraph(
        "Para Hector's Tree Service, la recomendación principal es el paquete Growth de $995 al mes. Ofrece el mejor equilibrio entre producción visible de marketing, mantenimiento del sitio y valor operativo sin comprometer demasiado presupuesto al inicio."
    )
    doc.add_paragraph(
        "Es un paquete suficientemente fuerte para mantener activa la marca con contenido constante, reforzar presencia en Google y redes sociales, y sostener visibilidad repetida con una inversión razonable para un negocio local de servicios."
    )

    doc.add_paragraph("Aprobación y Próximos Pasos", style="Heading 1")
    doc.add_paragraph(
        "Si se aprueba la propuesta, el proyecto puede avanzar con soporte de lanzamiento del sitio, programación mensual de contenido y planeación inicial de onboarding para DocSetter. Esta propuesta es válida por 30 días a partir de la fecha indicada."
    )

    add_on_table = doc.add_table(rows=1, cols=2)
    add_on_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    add_on_table.autofit = False
    set_table_layout(add_on_table)
    add_on_table.columns[0].width = Inches(4.4)
    add_on_table.columns[1].width = Inches(1.8)
    for cell, text in zip(add_on_table.rows[0].cells, ["Add-On Opcional", "Tarifa"]):
        set_cell_shading(cell, "E7F2EB")
        set_cell_border(cell, color="B7CFBC", size="10")
        set_cell_margins(cell)
        r = cell.paragraphs[0].add_run(text)
        r.font.bold = True
        r.font.color.rgb = ACCENT_DARK
    for name, rate in [
        ("Video corto adicional editado", "$125 c/u"),
        ("Sesión adicional con dron", "$250 c/u"),
        ("Landing page o página de servicio adicional", "$225 c/u"),
    ]:
        row = add_on_table.add_row()
        for cell in row.cells:
            set_cell_border(cell)
            set_cell_margins(cell)
        row.cells[0].paragraphs[0].add_run(name)
        row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        row.cells[1].paragraphs[0].add_run(rate)

    terms = doc.add_paragraph()
    terms.add_run("Nota de términos: ").font.bold = True
    terms.add_run(
        "los paquetes mensuales están pensados para mantener consistencia. Las grabaciones en sitio y con dron dependen del clima, acceso a la propiedad y condiciones seguras de vuelo."
    )

    prepared = doc.add_paragraph()
    prepared.add_run("Preparado por: ").font.bold = True
    prepared.add_run("Israel Aguilar | CEO, DestinyDevelopment")

    sign = doc.add_table(rows=2, cols=2)
    sign.alignment = WD_TABLE_ALIGNMENT.LEFT
    sign.autofit = False
    set_table_layout(sign)
    sign.columns[0].width = Inches(3.1)
    sign.columns[1].width = Inches(3.1)
    labels = [("Aprobación del Cliente", "Fecha"), ("Hector Aguilar", "__________________________")]
    for row_idx, (left_text, right_text) in enumerate(labels):
        row = sign.rows[row_idx]
        for cell in row.cells:
            set_cell_border(cell, color="FFFFFF", size="0")
            set_cell_margins(cell, top=110, bottom=110)
        row.cells[0].paragraphs[0].add_run(left_text).font.bold = row_idx == 0
        row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        row.cells[1].paragraphs[0].add_run(right_text).font.bold = row_idx == 0


def build_document() -> Path:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    apply_styles(doc)
    setup_page(doc)
    add_cover(doc)
    add_body(doc)
    doc.save(OUTPUT_PATH)
    return OUTPUT_PATH


if __name__ == "__main__":
    path = build_document()
    print(path)
